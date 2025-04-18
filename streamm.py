from docx import Document as DocxDocument
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains.retrieval_qa.base import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.schema import Document
import tempfile
import dotenv
import os

dotenv.load_dotenv()
api_key   = os.getenv("api_key")
if "chat_history" not in st.session_state:
    st.session_state.chat_history = ""

if "query_input" not in st.session_state:
    st.session_state.query_input = ""

if "docs" not in st.session_state:
    st.session_state.docs = []

if "qa" not in st.session_state:
    st.session_state.qa = None

if "clear_chat_clicked" not in st.session_state:
     st.session_state.clear_chat_clicked = False 

if not api_key:
     st.warning("API key not found , please check env variables")
else:    
    if "llm" not in st.session_state:    
        st.session_state.llm = ChatOpenAI(
        base_url  = "https://api.groq.com/openai/v1",
        api_key   = api_key,
      #  model_name= "llama3-70b-8192"  #model name reached limit
        model_name= "llama3-8b-8192"
    )   
        
def clear_text():
    st.session_state.query=st.session_state.query_input
    st.session_state.query_input=""

def buildResult(response):
                    st.write("You : "+response["query"])   
                    st.write("HRBuddy : "+response["result"])   
                    #put/append into chat box                                                 
                    st.session_state.chat_history +="\tYou:\n"+response["query"] + "\n"
                    line = '_' * 80               
                    st.session_state.chat_history +="\n\t\t\t\t\t\t\t\t\t\t\t\t\t\t\tHRBuddy:\n"
                    st.session_state.chat_history +=response["result"]+"\n"                  
                    st.session_state.chat_history+=line + "\n"
                    # st.text_area("Conversation with HRBuddy", value=st.session_state.chat_history, height=400)
                    st.write("Full Conversation with HRBuddy:")
                    render_chat(st.session_state.chat_history)


def start_chat(qa):
    if  st.session_state.clear_chat_clicked :
         return
    response = qa.invoke("Is this a resume or is resume uploaded, answer in yes or no.")
    ans = response["result"].lower()
    if "yes" not in ans:
         st.error("The candidate's reusme is not uploaded! Please upload resume")
    else: 
        query = st.session_state.get("query",'')
        upQuery = query.upper()
        
        if upQuery == "Q" or upQuery == "QUIT" or upQuery == "END" or upQuery == "BYE" :
            st.write("Thank you for using our AI system, Good bye..!")
        else:
            if query != "" :
                with st.spinner("Analyzing data based on your query.."):
                    response = qa.invoke(query)
                    #if not response["source_documents"]:
                        #st.write("The Question is out of context or the correct documents not provided")
                    #else:
                    # if query != "":
                    buildResult(response)
                    query=""
            

def render_chat(chat_history: str):
    st.markdown("""
    <style>
    .chat-box {
        max-height: 400px;
        overflow-y: auto;
        color: #784212;
        border: 1px solid #444;
        border-radius: 5px;   
  
    }
    </style>
    """, unsafe_allow_html=True)
    st.markdown(f"<div class='chat-box' id='chatbox'>{chat_history}</div>", unsafe_allow_html=True)


def load_file(file_path):
    docx  = DocxDocument(file_path)
    text  = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
    return [Document(page_content=text)]

def splitData(pdfData):
   splitter = RecursiveCharacterTextSplitter(chunk_size=500,chunk_overlap=100)
   splittedpdfDoc = splitter.split_documents(pdfData)
   return splittedpdfDoc


def loadAndSplit_pdfFile(files):  
    loader  = PyMuPDFLoader(files)
    pdfData = loader.load()
    splittedpdfDoc = splitData(pdfData)
    return splittedpdfDoc

def createTempfile(files):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(files.read())
        return tmp_file.name

def  createRetrieverQA(docs,emb):
    vectorStores = FAISS.from_documents(docs,emb)
    retriever = vectorStores.as_retriever()
    qa = RetrievalQA.from_llm(
    llm = st.session_state.llm,
    retriever = retriever,
    return_source_documents=False       
    )
    return qa


if "emb" not in st.session_state:
    st.session_state.emb= HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

st.title("OnBoarding Pro")
st.write("Dear HR, How can I help you today?")
st.text_input("Enter your query here", key="query_input", on_change=clear_text)

#st.header("******Your Personalized Hiring assistant******")
path = st.sidebar.file_uploader("Upload Candidate's Resume and supporting docs(*.pdf, *.docx)",
accept_multiple_files = True,
help="Upload one more more files"
)
if st.sidebar.button("Clear Chat",):
    st.text_area("Conversation with HRBuddy", value="", height=400)
    # render_chat("")
    st.session_state.clear_chat_clicked = True
    st.session_state.chat_history = ""   
    #   st.session_state.query_input="" 
 #   query=""                   
else:          
    st.session_state.clear_chat_clicked = False


if path is not None:
    new_docs = []
    for files in path:
        tmp = createTempfile(files)
        if files.name.endswith(".pdf"):         
            splittedpdfDoc = loadAndSplit_pdfFile(tmp)
            new_docs.extend(splittedpdfDoc)
        else:
            new_docs.extend(load_file(tmp))
  
if len(new_docs) > 0:
    st.session_state.docs = new_docs
    with st.spinner("Builing dataset, please wait.."):       
        st.session_state.qa =  createRetrieverQA(new_docs,st.session_state.emb)   
if st.session_state.qa: 
        start_chat(st.session_state.qa)
       

